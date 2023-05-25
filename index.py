import pandas as pd
from docx import Document
from docxcompose.composer import Composer
import os.path


directory = input("The directory where the files are located: ")
# C:/Users/User/Desktop/Xlsx-docx

sheetName = input("Sheet name: ")
# Vente

colName = input("Sheet column name(s): ")
# A, B (or only A)

docName = input("Document name(s): ")
# A, B (or only A)

mainDocName = input("Main document name: ")
# Modele Ordre du jour


def updateWord(sheetName, colName, docName, mainDocName):

    colNames = colName.replace(" ", "").split(",")
    docNames = docName.replace(" ", "").split(",")
    listVals = []
    slaves = []

    xls = pd.ExcelFile(f"{directory}/Information.xlsx")
    sheet = pd.read_excel(xls, sheet_name=sheetName)
    for col in sheet.columns:
        for column in colNames:
            if col == column:
                data = sheet.iloc[:, sheet.columns.get_loc(col) + 1].tolist()
                for value in data:
                    if str(value) != "nan":
                        listVals.append(value)
                doc = Document(f"{directory}/{docNames[0]}.docx")
                if len(listVals) != 0:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    words = p.text.split(' ')
                                    for elem in range(len(words)):
                                        if len(listVals) != 0:
                                            if "____" in words[elem]:
                                                words[elem] = words[elem].replace("____", str(listVals[0]))
                                                listVals.remove(listVals[0])
                                            elif "____" == words[elem]:
                                                words[elem] = words[elem].replace("____", str(listVals[0]))
                                                listVals.remove(listVals[0])
                                            else:
                                                words[elem] = words[elem]
                                        p.text = ' '.join(words)
                doc.save(f"{directory}/New_{docNames[0]}.docx")
                slaves.append(Document(f"{directory}/New_{docNames[0]}.docx"))
                docNames.remove(docNames[0])
            listVals = []

    def creatingDocumentName(mainDocName):
        newMainDocName = None
        if "." in mainDocName:
            array = mainDocName.split(".")
            print()
            if len(array) > 1:
                num = int(array[0]) + 1
                newMainDocName = f"{str(num)}.{array[1]}"
                if os.path.isfile(f"{directory}/{str(num)}.{array[1]}.docx"):
                    newMainDocName = f"{str(num) + 1}.{array[1]}"
        else:
            newMainDocName = f"0.{str(mainDocName)}"
            if os.path.isfile(f"{directory}/0.{str(mainDocName)}.docx"):
                newMainDocName = f"1.{str(mainDocName)}"
        return newMainDocName

    master = Document(f"{directory}/{mainDocName}.docx")
    composer = Composer(master)
    for slave in slaves:
        composer.append(slave)
    composer.save(f"{directory}/{creatingDocumentName(mainDocName)}.docx")


updateWord(sheetName, colName, docName, mainDocName)
